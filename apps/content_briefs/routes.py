from flask import Blueprint, render_template, request, jsonify, send_file, abort
import os
import json
import io
from apps.content_briefs.config import Config
from extensions import csrf
from apps.content_briefs.tasks.generate_brief import generate_brief_task
try:
    from docx import Document
except ImportError:
    Document = None

content_briefs_bp = Blueprint(
    'content_briefs_bp', __name__,
    template_folder='templates',
    static_folder='static'
)

BRIEFS_DIR = Config.BRIEFS_DIR

@content_briefs_bp.route('/', methods=['GET'])
def index():
    return render_template('briefs_index.html')

@content_briefs_bp.route('/start', methods=['POST'])
def start_brief():
    if request.is_json:
        data = request.get_json()
    else:
        data = request.form
    keyword = data.get('keyword')
    website = data.get('website')
    if not keyword or not website:
        return jsonify({'error': 'Missing keyword or website'}), 400
    task = generate_brief_task.apply_async(args=[keyword, website])
    return jsonify({'task_id': task.id}), 202

@content_briefs_bp.route('/download/<task_id>')
def download_brief(task_id):
    path = os.path.join(BRIEFS_DIR, f'brief_{task_id}.json')
    try:
        return send_file(path, as_attachment=True)
    except FileNotFoundError:
        abort(404)

@content_briefs_bp.route('/download_docx/<task_id>')
def download_docx(task_id):
    if Document is None:
        return "python-docx is not installed.", 500
    path = os.path.join(BRIEFS_DIR, f'brief_{task_id}.json')
    try:
        with open(path, 'r') as f:
            data = json.load(f)
        brief = data.get('brief', {})
        doc = Document()
        if brief.get('title'):
            doc.add_heading(brief['title'], 0)
        if brief.get('introduction'):
            doc.add_paragraph('Introduction: ' + brief['introduction'])
        if brief.get('audience'):
            doc.add_paragraph('Audience: ' + brief['audience'])
        if brief.get('search_intent'):
            doc.add_paragraph('Search Intent: ' + brief['search_intent'])
        if brief.get('talking_points'):
            doc.add_paragraph('Main Talking Points:')
            for tp in brief['talking_points']:
                doc.add_paragraph(tp, style='List Bullet')
        if brief.get('word_count'):
            doc.add_paragraph('Suggested Word Count: ' + str(brief['word_count']))
        if brief.get('full_brief'):
            doc.add_paragraph(brief['full_brief'])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name='blog_brief.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"DOCX download error: {e}")
        return "Brief not found or could not generate DOCX.", 404

@content_briefs_bp.route('/results/<task_id>')
def results_page(task_id):
    try:
        # Load the brief data
        brief_path = os.path.join(BRIEFS_DIR, f'brief_{task_id}.json')
        research_path = os.path.join(BRIEFS_DIR, f'research_{task_id}.json')
        
        content_brief = {}
        research_data = {}
        
        # Load brief data
        if os.path.exists(brief_path):
            with open(brief_path, 'r') as f:
                content_brief = json.load(f)
        
        # Load research data
        if os.path.exists(research_path):
            with open(research_path, 'r') as f:
                research_data = json.load(f)
        
        return render_template('brief_results.html',
                             content_brief=content_brief,
                             research_data=research_data,
                             task_id=task_id)
    except Exception as e:
        return f"Error loading brief: {e}", 404

@content_briefs_bp.route('/admin')
def admin_page():
    briefs = []
    try:
        for fname in os.listdir(BRIEFS_DIR):
            if fname.startswith('brief_') and fname.endswith('.json'):
                task_id = fname.replace('brief_', '').replace('.json', '')
                brief_path = os.path.join(BRIEFS_DIR, fname)
                research_path = os.path.join(BRIEFS_DIR, f'research_{task_id}.json')
                
                # Load brief data
                content_brief = {}
                research_data = {}
                
                if os.path.exists(brief_path):
                    with open(brief_path, 'r') as f:
                        content_brief = json.load(f)
                
                if os.path.exists(research_path):
                    with open(research_path, 'r') as f:
                        research_data = json.load(f)
                
                # Get topic from brief title or research keywords
                topic = 'Untitled'
                if content_brief.get('title', {}).get('main_title'):
                    topic = content_brief['title']['main_title']
                elif research_data.get('all_keywords'):
                    topic = research_data['all_keywords'][0] if research_data['all_keywords'] else 'Untitled'
                
                # Get website from research data
                website = research_data.get('website_url', 'N/A')
                if website and len(website) > 50:
                    website = website[:50] + '...'
                
                # Get creation time
                creation_time = os.path.getctime(brief_path)
                from datetime import datetime
                created_at = datetime.fromtimestamp(creation_time).strftime('%Y-%m-%d %H:%M')
                
                briefs.append({
                    'task_id': task_id,
                    'topic': topic,
                    'website': website,
                    'created_at': created_at,
                    'json_link': f'/content-briefs/download/{task_id}',
                    'docx_link': f'/content-briefs/download_docx/{task_id}',
                    'results_link': f'/content-briefs/results/{task_id}',
                    'has_research': os.path.exists(research_path)
                })
        
        # Sort by creation time (newest first)
        briefs.sort(key=lambda x: x['created_at'], reverse=True)
        
    except Exception as e:
        return f"Error loading briefs: {e}", 500
    return render_template('briefs_admin.html', briefs=briefs)


@content_briefs_bp.route('/progress/<task_id>')
def task_progress(task_id):
    task = generate_brief_task.AsyncResult(task_id)
    if task.state == 'PENDING':
        response = {'state': task.state, 'progress': 0, 'message': 'Task pending...'}
    elif task.state == 'PROGRESS':
        response = {'state': task.state, 'progress': task.info.get('step', 0), 'message': task.info.get('message', '')}
    elif task.state == 'SUCCESS':
        response = {'state': task.state, 'result': task.result.get('result', '')}
    else:
        response = {'state': task.state, 'message': str(task.info)}
    return jsonify(response)

def format_reddit_summary(text):
    import re
    if not text:
        return ''
    lines = re.split(r'(?<=\n)', text)
    items = []
    buffer = ''
    for line in lines:
        m = re.match(r'\s*(\d+)\.\s+(.*)', line)
        if m:
            if buffer:
                items.append({'type': 'p', 'text': buffer.strip()})
                buffer = ''
            items.append({'type': 'li', 'text': m.group(2).strip()})
        else:
            buffer += line
    if buffer.strip():
        items.append({'type': 'p', 'text': buffer.strip()})
    html = ''
    li_items = [i['text'] for i in items if i['type'] == 'li']
    if li_items:
        html += '<ol>' + ''.join(f'<li>{x}</li>' for x in li_items) + '</ol>'
    for i in items:
        if i['type'] == 'p' and i['text']:
            html += f'<p>{i["text"]}</p>'
    return html 

@content_briefs_bp.route('/test-form', methods=['GET', 'POST'])
def test_form():
    if request.method == 'POST':
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form
        keyword = data.get('keyword')
        website = data.get('website')
    return render_template('test_form.html') 